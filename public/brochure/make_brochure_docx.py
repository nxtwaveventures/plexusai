"""
Generates two styled PlexusAI brochure DOCX files:
  plexusai-brochure.docx        — regular (11pt body)
  plexusai-brochure-large.docx  — large print (14pt body)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────
TEAL        = RGBColor(0x0D, 0x94, 0x88)
TEAL_LIGHT  = RGBColor(0xF0, 0xFD, 0xFA)
NAVY        = RGBColor(0x1E, 0x29, 0x3B)
NAVY_DARK   = RGBColor(0x0F, 0x17, 0x2A)
BODY        = RGBColor(0x47, 0x55, 0x69)
MUTED       = RGBColor(0x94, 0xA3, 0xB8)
BORDER_COL  = RGBColor(0xE2, 0xE8, 0xF0)
GREEN       = RGBColor(0x05, 0x96, 0x69)
BLUE        = RGBColor(0x25, 0x63, 0xEB)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
OFF_WHITE   = RGBColor(0xF8, 0xFA, 0xFC)


def set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell border colours/sizes."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def set_page_margins(doc, top=1.5, bottom=1.5, left=1.8, right=1.8):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)
    section.page_width    = Cm(21)   # A4
    section.page_height   = Cm(29.7)


def add_run(para, text, bold=False, italic=False, size=None, color=None, font='DM Sans'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def heading(doc, text, level=1, size=22, color=NAVY, space_before=18, space_after=8, accent_word=None):
    """Add a heading paragraph; optionally colour one word in teal."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(para, before=space_before*20, after=space_after*20)
    if accent_word and accent_word in text:
        parts = text.split(accent_word, 1)
        add_run(para, parts[0], bold=True, size=size, color=color)
        add_run(para, accent_word, bold=True, size=size, color=TEAL)
        if parts[1]:
            add_run(para, parts[1], bold=True, size=size, color=color)
    else:
        add_run(para, text, bold=True, size=size, color=color)
    return para


def eyebrow(doc, text, size=7, space_before=14, space_after=4):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=space_before*20, after=space_after*20)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.name = 'DM Sans'
    run.font.size = Pt(size)
    run.font.color.rgb = TEAL
    return para


def body_para(doc, text, size=10, color=BODY, space_before=0, space_after=6, max_width=None):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=space_before*20, after=space_after*20, line=276)
    add_run(para, text, size=size, color=color)
    return para


def divider(doc):
    para = doc.add_paragraph()
    set_paragraph_spacing(para, before=10*20, after=10*20)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break(doc):
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def callout_box(doc, text, body_size=10):
    """Teal left-border callout using a 2-column table (thin accent | text)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    # Remove outer borders
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))

    w_left = table.columns[0].width = Cm(0.25)
    table.columns[1].width = Cm(14)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    set_cell_bg(left_cell, TEAL)
    set_cell_bg(right_cell, TEAL_LIGHT)

    p = right_cell.paragraphs[0]
    set_paragraph_spacing(p, before=8*20, after=8*20, line=276)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.name = 'DM Sans'
    run.font.size = Pt(body_size + 0.5)
    run.font.color.rgb = RGBColor(0x0a, 0x3d, 0x38)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def pipeline_table(doc, steps, body_size=10):
    """4-step pipeline table."""
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    col_w = Cm(3.7)
    for i, col in enumerate(table.columns):
        col.width = col_w

    labels = ['Onboard', 'Deploy', 'Validate', 'Certify']
    descs  = [
        'Intake, scope & IEC-aligned partner alignment',
        'Real OPD / IPD / ICU integration under oversight',
        'Clinician feedback & peer-reviewed RWE capture',
        'CDSCO-ready regulatory dossier & publication',
    ]

    for i in range(4):
        num_cell  = table.cell(0, i)
        desc_cell = table.cell(1, i)

        bg = TEAL if i == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(num_cell, bg)
        set_cell_bg(desc_cell, bg)

        # Number
        p = num_cell.paragraphs[0]
        set_paragraph_spacing(p, before=6*20, after=2*20)
        r = p.add_run(f'0{i+1}')
        r.bold = True
        r.font.name = 'DM Sans'
        r.font.size = Pt(body_size + 6)
        r.font.color.rgb = WHITE if i == 0 else BORDER_COL

        # Label
        p2 = num_cell.add_paragraph()
        set_paragraph_spacing(p2, before=0, after=4*20)
        r2 = p2.add_run(labels[i])
        r2.bold = True
        r2.font.name = 'DM Sans'
        r2.font.size = Pt(body_size + 1)
        r2.font.color.rgb = WHITE if i == 0 else NAVY

        # Description
        p3 = desc_cell.paragraphs[0]
        set_paragraph_spacing(p3, before=4*20, after=8*20, line=240)
        r3 = p3.add_run(descs[i])
        r3.font.name = 'DM Sans'
        r3.font.size = Pt(body_size - 1)
        r3.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0) if i == 0 else MUTED


def focus_row(doc, num, tag, tag_color, title, desc, body_size=10):
    """One focus area row: left accent panel | right content."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(10.8)

    left  = table.cell(0, 0)
    right = table.cell(0, 1)

    # Left bg tint
    r, g, b = tag_color[0], tag_color[1], tag_color[2]
    tint = RGBColor(
        min(255, r + int((255 - r) * 0.93)),
        min(255, g + int((255 - g) * 0.93)),
        min(255, b + int((255 - b) * 0.93)),
    )
    set_cell_bg(left, tint)
    set_cell_bg(right, OFF_WHITE)
    left.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Decorative number
    pn = left.paragraphs[0]
    set_paragraph_spacing(pn, before=10*20, after=2*20)
    rn = pn.add_run(num)
    rn.bold = True
    rn.font.name = 'DM Sans'
    rn.font.size = Pt(body_size + 14)
    rn.font.color.rgb = RGBColor(
        min(255, r + int((255-r)*0.80)),
        min(255, g + int((255-g)*0.80)),
        min(255, b + int((255-b)*0.80)),
    )

    # Category tag
    pt = left.add_paragraph()
    set_paragraph_spacing(pt, before=0, after=10*20)
    rt = pt.add_run(tag.upper())
    rt.bold = True
    rt.font.name = 'DM Sans'
    rt.font.size = Pt(7)
    rt.font.color.rgb = tag_color

    # Right: title + desc
    ph = right.paragraphs[0]
    set_paragraph_spacing(ph, before=10*20, after=4*20)
    rh = ph.add_run(title)
    rh.bold = True
    rh.font.name = 'DM Sans'
    rh.font.size = Pt(body_size + 1)
    rh.font.color.rgb = NAVY

    pd = right.add_paragraph()
    set_paragraph_spacing(pd, before=0, after=10*20, line=264)
    rd = pd.add_run(desc)
    rd.font.name = 'DM Sans'
    rd.font.size = Pt(body_size)
    rd.font.color.rgb = BODY


def program_row(doc, num, badge, badge_color, title, tagline, desc, cta, cta_color=TEAL, flagship=False, body_size=10):
    """One program row."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    table.columns[0].width = Cm(1.1)
    table.columns[1].width = Cm(11.6)
    table.columns[2].width = Cm(2.3)

    num_cell  = table.cell(0, 0)
    main_cell = table.cell(0, 1)
    cta_cell  = table.cell(0, 2)

    num_cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
    main_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cta_cell.vertical_alignment  = WD_ALIGN_VERTICAL.TOP

    # Number
    pnum = num_cell.paragraphs[0]
    set_paragraph_spacing(pnum, before=8*20, after=0)
    rnum = pnum.add_run(num)
    rnum.bold = True
    rnum.font.name = 'DM Sans'
    rnum.font.size = Pt(body_size + 4)
    rnum.font.color.rgb = TEAL if flagship else BORDER_COL

    # Badge
    pb = main_cell.paragraphs[0]
    set_paragraph_spacing(pb, before=8*20, after=3*20)
    rb = pb.add_run(f'  {badge.upper()}  ')
    rb.bold = True
    rb.font.name = 'DM Sans'
    rb.font.size = Pt(6.5)
    rb.font.color.rgb = badge_color
    if flagship:
        rb2 = pb.add_run('  MOST IMPACTFUL  ')
        rb2.bold = True
        rb2.font.name = 'DM Sans'
        rb2.font.size = Pt(6.5)
        rb2.font.color.rgb = WHITE

    # Title
    pt = main_cell.add_paragraph()
    set_paragraph_spacing(pt, before=0, after=2*20)
    rt = pt.add_run(title)
    rt.bold = True
    rt.font.name = 'DM Sans'
    rt.font.size = Pt(body_size + 1)
    rt.font.color.rgb = NAVY

    # Tagline
    phl = main_cell.add_paragraph()
    set_paragraph_spacing(phl, before=0, after=3*20)
    rhl = phl.add_run(tagline)
    rhl.italic = True
    rhl.font.name = 'DM Sans'
    rhl.font.size = Pt(body_size - 0.5)
    rhl.font.color.rgb = badge_color

    # Description
    pd = main_cell.add_paragraph()
    set_paragraph_spacing(pd, before=0, after=8*20, line=252)
    rd = pd.add_run(desc)
    rd.font.name = 'DM Sans'
    rd.font.size = Pt(body_size - 0.5)
    rd.font.color.rgb = BODY

    # CTA
    pcta = cta_cell.paragraphs[0]
    set_paragraph_spacing(pcta, before=10*20, after=0)
    rcta = pcta.add_run(cta + ' →')
    rcta.bold = True
    rcta.font.name = 'DM Sans'
    rcta.font.size = Pt(7.5)
    rcta.font.color.rgb = cta_color


def feature_card(doc, num, title, desc, body_size=10):
    return (num, title, desc)   # collected and laid out in a 2-col table


def rai_box(doc, body_size=10):
    """Responsible AI commitment callout."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    table.columns[0].width = Cm(0.3)
    table.columns[1].width = Cm(14.7)

    set_cell_bg(table.cell(0, 0), TEAL)
    set_cell_bg(table.cell(0, 1), TEAL_LIGHT)

    right = table.cell(0, 1)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    pl = right.paragraphs[0]
    set_paragraph_spacing(pl, before=6*20, after=4*20)
    rl = pl.add_run('OUR RESPONSIBLE AI COMMITMENT')
    rl.bold = True
    rl.font.name = 'DM Sans'
    rl.font.size = Pt(6.5)
    rl.font.color.rgb = TEAL

    pd = right.add_paragraph()
    set_paragraph_spacing(pd, before=0, after=8*20, line=252)
    text = (
        'All AI tools validated through PlexusAI undergo Institutional Ethics Committee (IEC) review. '
        'Patient data is handled in compliance with India\'s Digital Personal Data Protection (DPDP) Act 2023. '
        'PlexusAI outputs are designed to assist — never replace — qualified healthcare professionals. '
        'All research is conducted under ICMR guidelines and applicable CDSCO regulatory frameworks.'
    )
    rd = pd.add_run(text)
    rd.font.name = 'DM Sans'
    rd.font.size = Pt(body_size - 1)
    rd.font.color.rgb = NAVY


def cta_band_block(doc, body_size=10):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(3.5)

    set_cell_bg(table.cell(0, 0), NAVY)
    set_cell_bg(table.cell(0, 1), TEAL)

    left  = table.cell(0, 0)
    right = table.cell(0, 1)
    left.vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    ph = left.paragraphs[0]
    set_paragraph_spacing(ph, before=10*20, after=4*20)
    rh = ph.add_run('Ready to Validate Your AI Responsibly?')
    rh.bold = True
    rh.font.name = 'DM Sans'
    rh.font.size = Pt(body_size + 3)
    rh.font.color.rgb = WHITE

    ps = left.add_paragraph()
    set_paragraph_spacing(ps, before=0, after=10*20)
    rs = ps.add_run('Apply to the Innovation Sandbox — accepting partners committed to safe, evidence-based clinical AI for Q2 2026.')
    rs.font.name = 'DM Sans'
    rs.font.size = Pt(body_size - 1)
    rs.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    pb = right.paragraphs[0]
    pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(pb, before=8*20, after=8*20)
    rb = pb.add_run('Apply Now →')
    rb.bold = True
    rb.font.name = 'DM Sans'
    rb.font.size = Pt(body_size)
    rb.font.color.rgb = WHITE


def features_2col(doc, features, body_size=10):
    """Lay out 4 features in a 2×2 grid table."""
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for col in table.columns:
        col.width = Cm(7.5)

    for idx, (num, title, desc) in enumerate(features):
        row_idx = idx // 2
        col_idx = idx % 2
        cell = table.cell(row_idx, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        pn = cell.paragraphs[0]
        set_paragraph_spacing(pn, before=10*20, after=6*20)
        rn = pn.add_run(num)
        rn.bold = True
        rn.font.name = 'DM Sans'
        rn.font.size = Pt(7)
        rn.font.color.rgb = MUTED

        pt = cell.add_paragraph()
        set_paragraph_spacing(pt, before=0, after=4*20)
        rt = pt.add_run(title)
        rt.bold = True
        rt.font.name = 'DM Sans'
        rt.font.size = Pt(body_size + 1)
        rt.font.color.rgb = NAVY

        pd = cell.add_paragraph()
        set_paragraph_spacing(pd, before=0, after=10*20, line=252)
        rd = pd.add_run(desc)
        rd.font.name = 'DM Sans'
        rd.font.size = Pt(body_size - 0.5)
        rd.font.color.rgb = BODY


# ══════════════════════════════════════════════════════════════
def build(output_path, body_size=10):
    doc = Document()
    set_page_margins(doc)

    # Remove default Normal style spacing
    style = doc.styles['Normal']
    style.font.name = 'DM Sans'
    style.font.size = Pt(body_size)

    # ────────────────── COVER ──────────────────
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4*20)
    r = p.add_run('PlexusAI')
    r.bold = True
    r.font.name = 'DM Sans'
    r.font.size = Pt(body_size + 5)
    r.font.color.rgb = NAVY
    r2 = p.add_run('   Company Overview 2026')
    r2.font.name = 'DM Sans'
    r2.font.size = Pt(7)
    r2.font.color.rgb = TEAL

    eyebrow(doc, "India's First Hospital-Embedded Clinical AI Validation Hub",
            size=7, space_before=30, space_after=6)

    heading(doc, 'We Test and Validate AI\nin Hospitals.',
            size=body_size + 16, space_before=6, space_after=10,
            accent_word='in Hospitals.')

    body_para(doc,
        'Bring your AI. We embed it in live hospital workflows, gather clinician-validated '
        'evidence, and certify it — with patient safety at every step.',
        size=body_size + 2, space_before=4, space_after=20)

    # Trust strip
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for col in table.columns:
        col.width = Cm(5)
    for i, txt in enumerate(['IEC-Approved Trials', 'DPDP Act 2023 Compliant', 'CDSCO Framework']):
        c = table.cell(0, i)
        set_cell_bg(c, TEAL_LIGHT)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, before=8*20, after=8*20)
        r = p.add_run(f'● {txt}')
        r.font.name = 'DM Sans'
        r.font.size = Pt(8)
        r.font.color.rgb = TEAL
        r.bold = True

    divider(doc)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4*20, after=4*20)
    r1 = p.add_run('support@plexusai.in')
    r1.bold = True; r1.font.name = 'DM Sans'
    r1.font.size = Pt(8); r1.font.color.rgb = NAVY
    r2 = p.add_run('   ·   Accepting partners for Q2 2026   ·   ')
    r2.font.name = 'DM Sans'; r2.font.size = Pt(8); r2.font.color.rgb = MUTED
    r3 = p.add_run('plexusai.in')
    r3.bold = True; r3.font.name = 'DM Sans'
    r3.font.size = Pt(8); r3.font.color.rgb = TEAL

    # ────────────────── PAGE 2 — MISSION & JOURNEY ──────────────────
    page_break(doc)

    eyebrow(doc, 'Who We Are', size=7, space_before=8, space_after=4)
    heading(doc, 'Accelerating Safe Clinical AI\nFrom Idea to Responsible Impact',
            size=body_size + 10, space_before=4, space_after=10,
            accent_word='Safe Clinical AI')

    body_para(doc,
        "PlexusAI is India's first hospital-embedded clinical AI validation hub — built to bridge "
        "the gap between promising AI tools and safe, proven clinical adoption. We partner with "
        "hospitals, AI startups, pharma companies, and clinicians to rigorously validate, ethically "
        "certify, and responsibly scale healthcare AI inside live workflows — with patient safety at every step.",
        size=body_size, space_before=0, space_after=12)

    callout_box(doc,
        "Most healthcare AI never reaches patients — not because the technology fails, but because it "
        "is never rigorously tested in real clinical environments. We exist to change that, responsibly.",
        body_size=body_size)

    divider(doc)

    eyebrow(doc, 'How It Works', size=7, space_before=14, space_after=4)
    heading(doc, 'The Validation Journey',
            size=body_size + 6, space_before=4, space_after=8)

    body_para(doc,
        'From first conversation to a published, certified outcome — every project follows a structured, '
        'ethics-governed four-stage process.',
        size=body_size, space_before=0, space_after=12)

    pipeline_table(doc, [], body_size=body_size)

    # ────────────────── PAGE 3 — FOCUS AREAS ──────────────────
    page_break(doc)

    eyebrow(doc, 'What We Work On', size=7, space_before=8, space_after=4)
    heading(doc, 'Three Pillars of Hospital AI',
            size=body_size + 10, space_before=4, space_after=10,
            accent_word='Hospital AI')

    body_para(doc,
        'We operate across three areas of hospital AI — from daily clinical workflows '
        'to regulatory approval and published real-world evidence.',
        size=body_size, space_before=0, space_after=16)

    focus_row(doc, '01', 'Clinical Intelligence', TEAL,
              'AI Clinical Workflow',
              'We deploy AI to support discharge summaries, outpatient notes, ICU alerts, and clinical '
              'decision-making — reducing cognitive load on clinicians and improving patient outcomes. '
              'AI assists; clinicians decide.',
              body_size=body_size)

    doc.add_paragraph()  # spacer

    focus_row(doc, '02', 'Research & Trials', BLUE,
              'Digital Health Trials',
              'We run ethics-approved, structured trials for wearables, AI diagnostics, and remote monitoring '
              '— under Institutional Ethics Committee (IEC) oversight. You receive publishable, '
              'peer-reviewable real-world evidence.',
              body_size=body_size)

    doc.add_paragraph()

    focus_row(doc, '03', 'Regulatory Validation', GREEN,
              'Evidence Generation (RWE)',
              'Auditable, peer-reviewable real-world evidence for CDSCO regulatory submissions, '
              'investor due diligence, and multi-centre hospital adoption — fully compliant '
              'with DPDP Act 2023.',
              body_size=body_size)

    # ────────────────── PAGE 4 — PROGRAMS ──────────────────
    page_break(doc)

    eyebrow(doc, 'Programs Open for Application', size=7, space_before=8, space_after=4)
    heading(doc, 'What Can You Apply For?',
            size=body_size + 10, space_before=4, space_after=10,
            accent_word='Apply For?')

    body_para(doc,
        'Whether you are a startup seeking rigorous clinical proof, a pharma company funding responsible '
        'hospital AI research, or a clinician leading ethical digital transformation — there is a track for you.',
        size=body_size, space_before=0, space_after=16)

    divider(doc)
    programs = [
        ('01', 'For Startups',          TEAL,  False,
         'Clinical Validation Access',
         'Test your AI inside a real hospital — not a lab.',
         'Structured, ethics-overseen access to live hospital workflows, appropriately de-identified clinical '
         'data, and direct clinician feedback. Compress years of validation into months — without cutting corners.',
         'Apply for Access', TEAL),
        ('02', 'Flagship Program',      TEAL,  True,
         'PlexusAI Validation Certificate',
         'The credential that opens regulated markets for your AI.',
         "India's first structured clinical AI validation certificate. Investors gain confidence. "
         'Regulators engage sooner. Hospitals adopt certified AI with institutional trust.',
         'Apply for Certification', TEAL),
        ('03', 'For Pharma & MedTech',  GREEN, False,
         'Sponsored Hospital Research',
         'Commission ethics-approved studies with measurable patient impact.',
         'Commission structured AI clinical studies inside partner hospitals — all conducted under IEC approval, '
         'ICMR guidelines, and CDSCO regulatory frameworks, producing publishable real-world evidence.',
         'Explore Research', GREEN),
        ('04', 'High-Potential Startups', BLUE, False,
         'AI Accelerator Track',
         'Skip the noise. Scale inside a hospital.',
         'Equity-based acceleration — co-building with clinicians, hospital administrators, and our expert '
         'advisory network to bring safe, validated AI to scale.',
         'Apply to Accelerator', BLUE),
        ('05', 'For Clinicians',        TEAL,  False,
         'AI Certification for Doctors',
         'Become the AI-fluent clinician your hospital needs.',
         'Hospital-grounded AI certification designed by clinicians, for clinicians. Evaluate, implement, '
         'and champion AI tools with clinical confidence and ethical clarity.',
         'Register', TEAL),
    ]
    for num, badge, badge_color, flagship, title, tagline, desc, cta, cta_color in programs:
        program_row(doc, num, badge, badge_color, title, tagline, desc, cta, cta_color, flagship, body_size)
        divider(doc)

    # ────────────────── PAGE 5 — WHY + CONTACT ──────────────────
    page_break(doc)

    eyebrow(doc, 'Why PlexusAI', size=7, space_before=8, space_after=4)
    heading(doc, 'What You Get With PlexusAI',
            size=body_size + 10, space_before=4, space_after=10,
            accent_word='With PlexusAI')

    body_para(doc,
        'We take your AI from concept to certified — inside real hospitals, with patient safety '
        'and regulatory compliance built into every step.',
        size=body_size, space_before=0, space_after=16)

    feats = [
        ('01', 'Validate in Real Hospitals',
         'We integrate your AI into live OPDs, IPDs, and ICUs — with real clinical data, under institutional ethics oversight. Validated in months, not years.'),
        ('02', 'Clinical Certification',
         'The credential that opens regulated doors. Investors gain confidence. Regulators engage sooner. Hospitals adopt certified AI with institutional trust.'),
        ('03', 'Ethically Sourced Data',
         'Appropriately de-identified patient data, governed under DPDP Act 2023 and IEC approval — enabling rigorous, compliant AI validation across live hospital workflows.'),
        ('04', 'Fast-Track to Market',
         'Direct access to clinicians, hospital administrators, and our advisory network — compressing the path from validated AI to responsibly adopted AI.'),
    ]
    features_2col(doc, feats, body_size=body_size)

    doc.add_paragraph()
    rai_box(doc, body_size=body_size)

    doc.add_paragraph()
    cta_band_block(doc, body_size=body_size)

    divider(doc)

    eyebrow(doc, 'Get In Touch', size=7, space_before=10, space_after=4)
    heading(doc, 'Contact Us', size=body_size + 4, space_before=4, space_after=8)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    for col in table.columns:
        col.width = Cm(7.5)

    for i, (lbl, val, link_color) in enumerate([
        ('Email', 'support@plexusai.in', TEAL),
        ('Website', 'plexusai.in', TEAL),
    ]):
        c = table.cell(0, i)
        set_cell_bg(c, OFF_WHITE)
        pl = c.paragraphs[0]
        set_paragraph_spacing(pl, before=8*20, after=3*20)
        rl = pl.add_run(lbl.upper())
        rl.bold = True; rl.font.name = 'DM Sans'
        rl.font.size = Pt(6.5); rl.font.color.rgb = MUTED
        pv = c.add_paragraph()
        set_paragraph_spacing(pv, before=0, after=8*20)
        rv = pv.add_run(val)
        rv.bold = True; rv.font.name = 'DM Sans'
        rv.font.size = Pt(body_size); rv.font.color.rgb = link_color

    divider(doc)
    pf = doc.add_paragraph()
    set_paragraph_spacing(pf, before=4*20, after=0)
    rf = pf.add_run('PlexusAI — Company Brochure 2026  ·  All rights reserved')
    rf.font.name = 'DM Sans'; rf.font.size = Pt(7); rf.font.color.rgb = MUTED

    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__)) + '/'
    build(base + 'plexusai-brochure.docx',       body_size=11)
    build(base + 'plexusai-brochure-large.docx',  body_size=14)
    print('Done.')
